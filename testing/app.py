from flask import Flask, render_template, request, send_from_directory
from docx import Document
import os
import re

app = Flask(__name__)

def preprocess_mcq_lines(lines):
    """
    Preprocesses the input lines to handle questions and options.
    """
    processed_lines = []
    current_question = []
    in_options = False

    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:  # Skip empty lines
            continue

        # Detect options with or without labels (e.g., (A), (B), A., B. etc.)
        option_match = re.match(r'^[\(\[]?[A-Za-z0-9][\)\]].*|^[A-Da-d]\..*', stripped_line)

        if option_match:
            if current_question and not in_options:  # Save previous question block
                processed_lines.append("\n".join(current_question).strip())
                current_question = []
            processed_lines.append(stripped_line)
            in_options = True
        else:
            # If the line does not match option patterns, consider it part of the question
            current_question.append(stripped_line)
            in_options = False

    if current_question:  # Add the last question if any
        processed_lines.append("\n".join(current_question).strip())

    return processed_lines

def convert_text_to_word(lines, template_file):
    """
    Converts processed MCQ lines into a Word document format using a template file.
    """
    doc = Document(template_file)
    table_index = 0
    line_index = 0
    incorrect_options_lines = []

    lines = preprocess_mcq_lines(lines)

    while table_index < len(doc.tables) and line_index < len(lines):
        # Skip empty lines
        while line_index < len(lines) and not lines[line_index].strip():
            line_index += 1

        # Collect question lines
        question_lines = []
        while line_index < len(lines) and lines[line_index].strip() and not re.match(r'^[\(\[]?[A-Za-z0-9][\)\]].*|^[A-Da-d]\..*', lines[line_index]):
            question_lines.append(lines[line_index].strip())
            line_index += 1

        question = "\n".join(question_lines).strip()

        # Skip empty lines between question and options
        while line_index < len(lines) and not lines[line_index].strip():
            line_index += 1

        # Collect options and identify the correct answer
        options = []
        correct_index = None
        option_labels = ['A', 'B', 'C', 'D']  # Define the order of labels
        while line_index < len(lines) and lines[line_index].strip() and re.match(r'^[\(\[]?[A-Za-z0-9][\)\]].*|^[A-Da-d]\..*', lines[line_index]):
            option_line = lines[line_index].strip()
            
            # Extract the label and the option text
            label_match = re.match(r'^([\(\[]?)([A-Da-d])[\)\]].*(?:@)?', option_line)
            if label_match:
                label = label_match.group(2).upper()
                option_text = re.split(r'[\)\]].*', option_line, maxsplit=1)[-1].strip()
            else:
                label = f"{len(options)+1}"  # Fallback to number if label not found
                option_text = option_line

            # Check if this option is marked as correct
            if '@' in option_line:
                correct_index = len(options)  # 0-based index

            options.append((label, option_text))
            line_index += 1

        # If no correct answer is found, add line number to incorrect_options_lines
        if correct_index is None:
            incorrect_options_lines.append(line_index - len(options))

        # Insert data into the current table
        table = doc.tables[table_index]
        table.cell(0, 1).text = question
        table.cell(1, 1).text = 'multiple_choice'

        # Assign options to variables
        option1 = options[0][1] if len(options) > 0 else ''
        option2 = options[1][1] if len(options) > 1 else ''
        option3 = options[2][1] if len(options) > 2 else ''
        option4 = options[3][1] if len(options) > 3 else ''
        
        # Determine the correct answer number
        correct_number = (correct_index + 1) if correct_index is not None else ''

        # Place the options in the correct format
        table.cell(2, 0).text = 'Option'
        table.cell(2, 1).text = f'{option1}'  # Option 1
        table.cell(3, 0).text = f'{option2}'  # Option 2
        table.cell(3, 2).text = f'{option3}'  # Option 3
        table.cell(4, 1).text = f'{option4}'  # Option 4

        # Place the correct answer number in the answer row
        table.cell(4, 3).text = f'{correct_number}' if correct_number else ''

        # Move to the next table for the next question set
        table_index += 1

    return doc, incorrect_options_lines

@app.route('/', methods=['GET', 'POST'])
def index():
    """
    Handles the main index route for displaying the form and processing MCQ input.
    """
    if request.method == 'POST':
        text_content = request.form['text_content']
        template_size = request.form['template_size']
        lines = text_content.split('\n')

        template_file_path = get_template_path(template_size)

        if not template_file_path:
            return render_template('index.html', error="Invalid template size selected")

        doc, incorrect_options_lines = convert_text_to_word(lines, template_file_path)
        output_file_path = "static/output.docx"
        doc.save(output_file_path)

        return render_template('result.html', output_file="output.docx", incorrect_options_lines=incorrect_options_lines)

    return render_template('index.html')

def get_template_path(template_size):
    """
    Returns the file path of the template based on the selected size.
    """
    template_files = {
        '25': "templates/template 25 copy.docx",
        '50': "templates/template 50.docx",
        '100': "templates/template 100.docx",
        '150': "templates/template 150.docx"
    }
    return template_files.get(template_size)

@app.route('/download', methods=['GET'])
def download():
    """
    Allows users to download the generated Word document.
    """
    output_file_path = "static/output.docx"
    return send_from_directory(os.path.join(os.getcwd(), 'static'), 'output.docx', as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
