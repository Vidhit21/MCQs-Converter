from flask import Flask, render_template, request, send_from_directory
from docx import Document
import os
import re

app = Flask(__name__)

def split_text_by_pattern(lines):
    """
    Splits the input lines into chunks based on delimiters in the format ---FileName---.
    Extracts the filename from the delimiter and associates it with the corresponding text block.
    """
    results = []
    current_chunk = []
    current_filename = None

    for line in lines:
        stripped_line = line.strip()
        # Check if the line starts and ends with '---', indicating a new file
        if stripped_line.startswith('---') and stripped_line.endswith('---'):
            # Save the current chunk if it's not empty
            if current_chunk:
                results.append((current_filename, current_chunk))
                current_chunk = []
            # Extract filename from the delimiter
            current_filename = stripped_line.strip('-')
        else:
            # Add line to the current chunk
            current_chunk.append(stripped_line)
    
    # Add the last chunk if it exists
    if current_chunk:
        results.append((current_filename, current_chunk))
    
    return results

def preprocess_mcq_lines(lines):
    """
    Preprocesses the input lines to handle questions and options.
    """
    processed_lines = []
    current_question = []
    in_options = False

    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:  # Skip empty lines
            continue

        # Detect options with or without labels (e.g., (A), (B), A., B. etc.)
        option_match = re.match(r'^[\(\[]?[A-Za-z0-9][\)\]].*|^[A-Da-d]\..*', stripped_line)

        if option_match:
            if current_question and not in_options:  # Save previous question block
                processed_lines.append("\n".join(current_question).strip())
                current_question = []
            processed_lines.append(stripped_line)
            in_options = True
        else:
            # If the line does not match option patterns, consider it part of the question
            current_question.append(stripped_line)
            in_options = False

    if current_question:  # Add the last question if any
        processed_lines.append("\n".join(current_question).strip())

    return processed_lines

def convert_text_to_word(lines, template_file):
    """
    Converts processed MCQ lines into a Word document format using a template file.
    """
    doc = Document(template_file)
    table_index = 0
    line_index = 0
    incorrect_options_lines = []

    lines = preprocess_mcq_lines(lines)

    # Regex pattern to match question numbers like "Q.1", "Q1.", "1.", "Q1)" etc.
    question_number_pattern = re.compile(r'^(Q\.?\d+\.?|\d+\.?|\d+\))\s*')

    while table_index < len(doc.tables) and line_index < len(lines):
        # Skip empty lines
        while line_index < len(lines) and not lines[line_index].strip():
            line_index += 1

        # Collect question lines
        question_lines = []
        while line_index < len(lines) and lines[line_index].strip() and not re.match(r'^[\(\[]?[A-Za-z0-9][\)\]].*|^[A-Da-d]\..*', lines[line_index]):
            # Remove question number (e.g., "Q.1", "1.", "Q1)")
            cleaned_line = re.sub(question_number_pattern, '', lines[line_index].strip())
            question_lines.append(cleaned_line)
            line_index += 1

        question = "\n".join(question_lines).strip()

        # Skip empty lines between question and options
        while line_index < len(lines) and not lines[line_index].strip():
            line_index += 1

        # Collect options and identify the correct answer
        options = []
        correct_index = None
        while line_index < len(lines) and lines[line_index].strip() and re.match(r'^[\(\[]?[A-Za-z0-9][\)\]].*|^[A-Da-d]\..*', lines[line_index]):
            option_line = lines[line_index].strip()
            option = option_line.split(')', 1)[-1].strip()  # Get text after label
            options.append(option)
            if '@' in option_line:  # Identify correct answer
                correct_index = len(options) - 1
            line_index += 1

        # If no correct answer is found, add line number to incorrect_options_lines
        if correct_index is None:
            incorrect_options_lines.append(line_index - len(options))

        # Insert data into the current table
        table = doc.tables[table_index]
        table.cell(0, 1).text = question
        table.cell(1, 1).text = 'multiple_choice'

        # Assign options to variables
        option1 = options[0] if len(options) > 0 else ''
        option2 = options[1] if len(options) > 1 else ''
        option3 = options[2] if len(options) > 2 else ''
        option4 = options[3] if len(options) > 3 else ''

        # Place the options in the correct format
        table.cell(2, 0).text = 'Option'  
        table.cell(2, 1).text = f'{option1}'  # Option 1
        table.cell(3, 0).text = f'{option2}'  # Option 2
        table.cell(3, 2).text = f'{option3}'  # Option 3
        table.cell(4, 1).text = f'{option4}'  # Option 4

        # Place the correct answer index (1-based) in the answer row
        if correct_index is not None:
            # Since options are 0-indexed in Python, add 1 to get a 1-based index
            correct_option_index = correct_index + 1
            table.cell(4, 3).text = str(correct_option_index)
        else:
            table.cell(4, 3).text = ''

        # Move to the next table for the next question set
        table_index += 1

    return doc, incorrect_options_lines

@app.route('/', methods=['GET', 'POST'])
def index():
    """
    Handles the main index route for displaying the form and processing MCQ input.
    """
    if request.method == 'POST':
        # Get the text content or uploaded files
        text_content = request.form['text_content']
        lines = text_content.split('\n') if text_content else []

        # Split lines into chunks with filenames
        file_chunks = split_text_by_pattern(lines)

        # Template selection
        template_size = request.form['template_size']
        template_file_path = get_template_path(template_size)

        if not template_file_path:
            return render_template('index.html', error="Invalid template size selected")

        # Process each chunk and generate files
        file_names = []
        for i, (filename, chunk) in enumerate(file_chunks):
            try:
                doc, incorrect_options_lines = convert_text_to_word(chunk, template_file_path)

                # Use specified filename or default
                output_file_name = f"{filename or f'output_{i + 1}'}.docx"
                output_file_path = f"static/{output_file_name}"
                doc.save(output_file_path)
                file_names.append(output_file_name)
            except Exception as e:
                return render_template('index.html', error=f"Error processing chunk {i + 1}: {str(e)}")

        # Return results
        return render_template('result.html', file_names=file_names)

    return render_template('index.html')

def get_template_path(template_size):
    """
    Returns the file path of the template based on the selected size.
    """
    template_files = {
        '25': "templates/template 25.docx",
        '50': "templates/template 50.docx",
        '100': "templates/template 100.docx",
        '125': "templates/template 125.docx",
        '150': "templates/template 150.docx",
        '200': "templates/template 200.docx"
    }
    return template_files.get(template_size)

@app.route('/download/<filename>', methods=['GET'])
def download(filename):
    """
    Allows users to download the generated Word document.
    """
    return send_from_directory(os.path.join(os.getcwd(), 'static'), filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
